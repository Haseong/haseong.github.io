import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_markdown_text(text):
    """Clean markdown text for docx"""
    # Remove image references
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Convert bold markdown to plain text (will handle separately)
    # Convert italic markdown to plain text (will handle separately)
    # Remove inline code backticks
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

def add_formatted_paragraph(doc, text):
    """Add paragraph with basic markdown formatting preserved"""
    p = doc.add_paragraph()
    
    # Simple bold and italic handling
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            if part:
                p.add_run(part)
    
    return p

def convert_md_to_docx(md_file, docx_file):
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('AlphaCode Blog Posts Collection', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_section = []
    in_code_block = False
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if current_section:
                    code_text = ''.join(current_section)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                    current_section = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue
        
        if in_code_block:
            current_section.append(line)
            continue
        
        # Handle headers
        if line.startswith('## '):
            # Post title (H2)
            title_text = clean_markdown_text(line[3:].strip())
            doc.add_heading(title_text, level=1)
        elif line.startswith('### '):
            # Section title (H3)
            title_text = clean_markdown_text(line[4:].strip())
            doc.add_heading(title_text, level=2)
        elif line.startswith('#### '):
            # Subsection title (H4)
            title_text = clean_markdown_text(line[5:].strip())
            doc.add_heading(title_text, level=3)
        elif line.startswith('##### '):
            # Sub-subsection title (H5)
            title_text = clean_markdown_text(line[6:].strip())
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('# '):
            # Main title (H1) - skip as we already have it
            if not line.startswith('# AlphaCode'):
                title_text = clean_markdown_text(line[2:].strip())
                doc.add_heading(title_text, level=1)
        elif line.strip() == '---':
            # Horizontal rule - add page break
            doc.add_page_break()
        elif line.strip():
            # Regular paragraph
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                # Bullet list
                text = line.strip()[2:]
                p = doc.add_paragraph(clean_markdown_text(text))
                p.style = doc.styles['List Bullet']
            elif re.match(r'^\d+\.\s', line.strip()):
                # Numbered list
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(clean_markdown_text(text))
                p.style = doc.styles['List Number']
            elif line.strip().startswith('>'):
                # Blockquote
                text = line.strip().lstrip('> ')
                p = doc.add_paragraph()
                run = p.add_run(clean_markdown_text(text))
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            else:
                # Regular paragraph with formatting
                add_formatted_paragraph(doc, line.strip())
    
    # Save the document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")
    
    # Get file size
    import os
    size = os.path.getsize(docx_file)
    print(f"Output file size: {size/1024/1024:.2f} MB")

# Convert the combined markdown to docx
convert_md_to_docx('all_posts_combined.md', 'all_posts_combined.docx')
